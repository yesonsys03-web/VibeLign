# === ANCHOR: DOCX_RENDERER_START ===
from __future__ import annotations

import io

from vibelign.core.reporting_cli.font_sizes import ReportFontSizes
from vibelign.core.reporting_cli.fonts import ReportFonts, font_def
from vibelign.core.reporting_cli.models import ReportModel
from vibelign.core.reporting_cli.templates import meta_line
from vibelign.core.reporting_cli.themes import get_theme

try:
    import docx  # noqa: F401
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# === ANCHOR: DOCX_RENDERER_REPORTRENDERERUNAVAILABLE_START ===
class ReportRendererUnavailable(RuntimeError):
    """렌더러 라이브러리(python-docx/pptx) 미설치."""
# === ANCHOR: DOCX_RENDERER_REPORTRENDERERUNAVAILABLE_END ===


# === ANCHOR: DOCX_RENDERER__ADD_PAGE_NUMBER_FOOTER_START ===
def _add_page_number_footer(document) -> None:
    """바닥글 중앙에 'PAGE / NUMPAGES'(예: 1 / 3) 필드를 넣는다. Word 가 실제 번호 렌더."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    para = document.sections[0].footer.paragraphs[0]
    para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

    # === ANCHOR: DOCX_RENDERER__FIELD_START ===
    def _field(instr: str) -> None:
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr)
        para._p.append(fld)
    # === ANCHOR: DOCX_RENDERER__FIELD_END ===

    _field("PAGE")
# === ANCHOR: DOCX_RENDERER__ADD_PAGE_NUMBER_FOOTER_END ===
    para.add_run(" / ")
    _field("NUMPAGES")


# === ANCHOR: DOCX_RENDERER_RENDER_DOCX_START ===
def render_docx(
    model: ReportModel,
    theme: str = "classic",
    page_numbers: bool = False,
    font_sizes: ReportFontSizes | None = None,
    fonts: ReportFonts | None = None,
) -> bytes:
    if not DOCX_AVAILABLE:
        raise ReportRendererUnavailable(
            "Word 내보내기에 python-docx 가 필요합니다. (pip install python-docx)"
        )
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    accent = RGBColor.from_string(get_theme(theme).accent.lstrip("#"))
    heading_name = (
        font_def(fonts.heading).office_name if fonts and fonts.heading else None
    )
    body_name = font_def(fonts.body).office_name if fonts and fonts.body else None

    def _accent(heading) -> None:
        for r in heading.runs:
            r.font.color.rgb = accent

    def _size(paragraph, value: int | None) -> None:
        if value is None:
            return
        for run in paragraph.runs:
            run.font.size = Pt(value)

    def _font(paragraph, name: str | None) -> None:
        if name is None:
            return
        for run in paragraph.runs:
            run.font.name = name
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.insert(0, rfonts)
            for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rfonts.set(qn(attr), name)

    doc = Document()
    title = doc.add_heading(model.title, level=0)
    _accent(title)
    _size(title, font_sizes.title if font_sizes is not None else None)
    _font(title, heading_name)
    meta_para = doc.add_paragraph(meta_line(model))
    _size(meta_para, (font_sizes.meta or font_sizes.body) if font_sizes is not None else None)
    _font(meta_para, body_name)
    for section in model.sections:
        heading = doc.add_heading(section.heading, level=1)
        _accent(heading)
        _size(heading, font_sizes.heading if font_sizes is not None else None)
        _font(heading, heading_name)
        for block in section.blocks:
            if block.kind == "bullets":
                for item in block.items:
                    p = doc.add_paragraph(item, style="List Bullet")
                    _size(p, font_sizes.body if font_sizes is not None else None)
                    _font(p, body_name)
            elif block.kind == "summary":
                p = doc.add_paragraph()
                run = p.add_run(block.text)
                run.bold = True
                _size(p, font_sizes.body if font_sizes is not None else None)
                _font(p, body_name)
            else:  # paragraph
                p = doc.add_paragraph(block.text)
                _size(p, font_sizes.body if font_sizes is not None else None)
                _font(p, body_name)
    if page_numbers:
        _add_page_number_footer(doc)
    buf = io.BytesIO()
# === ANCHOR: DOCX_RENDERER_RENDER_DOCX_END ===
    doc.save(buf)
    return buf.getvalue()
# === ANCHOR: DOCX_RENDERER_END ===
