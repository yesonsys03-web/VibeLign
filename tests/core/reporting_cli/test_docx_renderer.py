import io
import pytest
from vibelign.core.reporting_cli.docx_renderer import render_docx, DOCX_AVAILABLE
from vibelign.core.reporting_cli.models import Block, ReportModel, Section


def _model():
    return ReportModel(
        title="예약 앱", report_type="work", date="2026-06-16",
        sections=[
            Section("개요", [Block(kind="summary", text="미용실 예약 앱")]),
            Section("핵심 내용", [Block(kind="bullets", items=["예약 캘린더", "알림 문자"])]),
            Section("배경", [Block(kind="paragraph", text="전화 예약 누락")]),
        ],
    )


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx 미설치")
def test_render_docx_opens_and_contains_text():
    data = render_docx(_model())
    assert data[:2] == b"PK"  # docx = zip
    from docx import Document
    doc = Document(io.BytesIO(data))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "예약 앱" in texts
    assert "업무 보고" in texts  # report_type 라벨
    assert "예약 캘린더" in texts
    assert "전화 예약 누락" in texts


def test_render_docx_missing_lib_raises_clear(monkeypatch):
    import vibelign.core.reporting_cli.docx_renderer as mod
    monkeypatch.setattr(mod, "DOCX_AVAILABLE", False)
    with pytest.raises(mod.ReportRendererUnavailable):
        render_docx(_model())


def test_docx_theme_applies_accent_to_heading():
    import io as _io

    import docx as _docx

    from vibelign.core.reporting_cli.docx_renderer import render_docx as _render
    from vibelign.core.reporting_cli.models import Block, ReportModel, Section

    m = ReportModel(title="t", report_type="work", date="d",
                    sections=[Section("개요", [Block(kind="summary", text="요약")])])
    data = _render(m, theme="executive")  # accent #1B3A6B
    d = _docx.Document(_io.BytesIO(data))
    colors = []
    for p in d.paragraphs:
        for r in p.runs:
            if r.font.color is not None and r.font.color.rgb is not None:
                colors.append(str(r.font.color.rgb))
    assert "1B3A6B" in colors


def test_docx_page_numbers_adds_page_field():
    import io as _io, docx as _docx
    from vibelign.core.reporting_cli.docx_renderer import render_docx as _r
    from vibelign.core.reporting_cli.models import Block, ReportModel, Section
    m = ReportModel(title="t", report_type="work", date="d",
                    sections=[Section("개요", [Block(kind="summary", text="x")])])
    data = _r(m, page_numbers=True)
    xml = _docx.Document(_io.BytesIO(data)).sections[0].footer._element.xml
    assert "PAGE" in xml and "NUMPAGES" in xml


def test_docx_no_page_numbers_no_footer_field():
    import io as _io, docx as _docx
    from vibelign.core.reporting_cli.docx_renderer import render_docx as _r
    from vibelign.core.reporting_cli.models import Block, ReportModel, Section
    m = ReportModel(title="t", report_type="work", date="d",
                    sections=[Section("개요", [Block(kind="summary", text="x")])])
    data = _r(m, page_numbers=False)
    xml = _docx.Document(_io.BytesIO(data)).sections[0].footer._element.xml
    assert "PAGE" not in xml
